"""
Procesador de documentos Word.
Utiliza python-docx para extraer texto.
"""
from pathlib import Path
from typing import Any

from docx import Document

from app.processors.base import Processor


class DocxProcessor(Processor):
    """
    Procesador para archivos Word (.docx).
    Extrae texto de párrafos y tablas.
    """
    supported_extensions = ["docx"]

    def process(self, file_path: Path) -> list[dict[str, Any]]:
        """
        Procesa un archivo Word y retorna una lista de fragmentos.

        Args:
            file_path: Ruta al archivo Word.

        Returns:
            Lista de diccionarios con 'content' y 'metadata'.
        """
        self.validate_file(file_path)

        try:
            doc = Document(file_path)
        except Exception as e:
            raise ValueError(f"Error al abrir archivo Word: {e}")

        chunks = []
        current_chunk = ""
        current_paragraphs = []
        chunk_count = 0

        # Extraer texto de párrafos
        for para in doc.paragraphs:
            text = para.text.strip()

            if not text:
                continue

            # Agregar al chunk actual
            current_paragraphs.append(text)
            current_chunk += text + "\n\n"

            # Si el chunk es suficientemente grande, guardarlo
            if len(current_chunk) > 1000:  # Aproximadamente 250 tokens
                chunk_count += 1
                chunk = {
                    "content": current_chunk.strip(),
                    "metadata": {
                        "type": "paragraphs",
                        "chunk_index": chunk_count,
                    },
                }
                chunks.append(chunk)
                current_chunk = ""
                current_paragraphs = []

        # Guardar el chunk restante
        if current_chunk.strip():
            chunk_count += 1
            chunk = {
                "content": current_chunk.strip(),
                "metadata": {
                    "type": "paragraphs",
                    "chunk_index": chunk_count,
                },
            }
            chunks.append(chunk)

        # Extraer texto de tablas
        for table_idx, table in enumerate(doc.tables):
            table_text = self._extract_table_text(table)

            if table_text.strip():
                chunk = {
                    "content": table_text,
                    "metadata": {
                        "type": "table",
                        "table_index": table_idx + 1,
                    },
                }
                chunks.append(chunk)

        if not chunks:
            raise ValueError("No se pudo extraer texto del documento Word")

        return chunks

    def _extract_table_text(self, table) -> str:
        """
        Extrae texto de una tabla.

        Args:
            table: Objeto tabla de python-docx.

        Returns:
            Texto extraído de la tabla.
        """
        parts = []

        for row in table.rows:
            row_parts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_parts.append(cell_text)
            if row_parts:
                parts.append(" | ".join(row_parts))

        return "\n".join(parts)
